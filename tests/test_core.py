"""Unit tests for DOCX Batch Updater core functionality.

This module contains tests for the document processing, format preservation,
and batch processing functionality.
"""

import unittest
import os
import tempfile
import shutil
from pathlib import Path

# Add src to path for imports
import sys
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'src'))

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    from core.docx_processor import DocxProcessor
    from core.batch_processor import BatchProcessor, ProcessingResult
    from src.utils.format_preserver import FormatPreserver
except ImportError as e:
    print(f"Import error: {e}")
    print("Please install dependencies: pip install -r requirements.txt")
    sys.exit(1)


class TestFormatPreserver(unittest.TestCase):
    """Test format preservation utilities."""

    def setUp(self):
        """Set up test fixtures."""
        self.preserver = FormatPreserver()

    def test_capture_and_apply_run_format(self):
        """Test capturing and applying run format."""
        # Create a test document with formatted text
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run("Test text")
        run.bold = True
        run.italic = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(255, 0, 0)

        # Capture format
        format_data = self.preserver.capture_run_format(run)

        # Verify captured data
        self.assertIn('bold', format_data)
        self.assertIn('italic', format_data)
        self.assertIn('font_size', format_data)
        self.assertIn('color_rgb', format_data)

    def test_capture_paragraph_format(self):
        """Test capturing paragraph format."""
        doc = Document()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)

        format_data = self.preserver.capture_paragraph_format(p)

        self.assertIn('alignment', format_data)
        self.assertIn('space_after', format_data)


class TestDocxProcessor(unittest.TestCase):
    """Test document processor functionality."""

    def setUp(self):
        """Set up test fixtures."""
        self.temp_dir = tempfile.mkdtemp()
        self.test_file = os.path.join(self.temp_dir, "test.docx")

        # Create a test document
        self._create_test_document()

    def tearDown(self):
        """Clean up test fixtures."""
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)

    def _create_test_document(self):
        """Create a test DOCX document."""
        doc = Document()

        # Add paragraph with text to replace
        p = doc.add_paragraph("Hello World 2024")

        # Add table
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "Cell 1 - 2024"
        table.rows[0].cells[1].text = "Cell 2 - 2024"
        table.rows[1].cells[0].text = "Cell 3 - 2024"
        table.rows[1].cells[1].text = "Cell 4 - 2024"

        # Add another paragraph
        p2 = doc.add_paragraph("More text 2024 here")

        doc.save(self.test_file)

    def test_load_document(self):
        """Test loading a document."""
        processor = DocxProcessor(self.test_file)
        self.assertTrue(processor.load())
        self.assertIsNotNone(processor.doc)

    def test_text_replacement(self):
        """Test text replacement in paragraphs."""
        processor = DocxProcessor(self.test_file)
        processor.load()

        count = processor.replace_text("2024", "2025")

        self.assertEqual(count, 6)  # 2 paragraphs + 4 table cells

        # Save and reload to verify
        processor.save(os.path.join(self.temp_dir, "output.docx"))
        processor2 = DocxProcessor(os.path.join(self.temp_dir, "output.docx"))
        processor2.load()

        # Check that replacement occurred
        text = processor2.doc.paragraphs[0].text
        self.assertIn("2025", text)
        self.assertNotIn("2024", text)

    def test_table_replacement(self):
        """Test text replacement in tables."""
        processor = DocxProcessor(self.test_file)
        processor.load()

        count = processor.replace_text("Cell 1 - 2024", "Cell 1 - UPDATED")

        self.assertEqual(count, 1)

        # Verify table cell was updated
        cell_text = processor.doc.tables[0].rows[0].cells[0].text
        self.assertEqual(cell_text, "Cell 1 - UPDATED")

    def test_backup_creation(self):
        """Test automatic backup creation."""
        processor = DocxProcessor(self.test_file)
        processor.load()

        backup_path = processor.create_backup()

        self.assertTrue(os.path.exists(backup_path))
        # New naming: {stem}_backup_{relhint}_{hash6}_{uuid8}.docx
        self.assertIn("_backup_", backup_path)
        self.assertTrue(backup_path.endswith(".docx"))

    def test_is_docx_file(self):
        """Test DOCX file validation."""
        # Valid DOCX file
        self.assertTrue(DocxProcessor.is_docx_file(self.test_file))

        # Invalid file
        invalid_file = os.path.join(self.temp_dir, "test.txt")
        with open(invalid_file, 'w') as f:
            f.write("Not a docx")
        self.assertFalse(DocxProcessor.is_docx_file(invalid_file))

    def test_get_statistics(self):
        """Test document statistics."""
        processor = DocxProcessor(self.test_file)
        processor.load()

        stats = processor.get_statistics()

        self.assertIn('paragraph_count', stats)
        self.assertIn('table_count', stats)
        self.assertIn('character_count', stats)
        self.assertIn('word_count', stats)
        self.assertIn('cell_count', stats)

        self.assertEqual(stats['paragraph_count'], 2)
        self.assertEqual(stats['table_count'], 1)
        self.assertEqual(stats['cell_count'], 4)


class TestBatchProcessor(unittest.TestCase):
    """Test batch processing functionality."""

    def setUp(self):
        """Set up test fixtures."""
        self.temp_dir = tempfile.mkdtemp()
        self.batch_processor = BatchProcessor(max_workers=2)

        # Create multiple test documents
        self.test_files = []
        for i in range(5):
            test_file = os.path.join(self.temp_dir, f"test_{i}.docx")
            self._create_test_document(test_file, f"Document {i} 2024")
            self.test_files.append(test_file)

    def tearDown(self):
        """Clean up test fixtures."""
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)

    def _create_test_document(self, file_path: str, content: str):
        """Create a test DOCX document."""
        doc = Document()
        p = doc.add_paragraph(content)
        p = doc.add_paragraph("Additional text 2024")
        doc.save(file_path)

    def test_process_documents(self):
        """Test processing multiple documents."""
        replacements = [("2024", "2025")]

        results = self.batch_processor.process_documents(
            self.test_files,
            replacements,
            create_backup=False
        )

        self.assertEqual(len(results), 5)
        self.assertTrue(all(r.success for r in results))

    def test_validate_files(self):
        """Test file validation."""
        # Add a non-existent file
        invalid_files = self.test_files + ["/nonexistent/file.docx"]

        valid_files = self.batch_processor._validate_files(invalid_files)

        self.assertEqual(len(valid_files), 5)

    def test_summary_statistics(self):
        """Test batch processing summary."""
        replacements = [("2024", "2025")]

        self.batch_processor.process_documents(
            self.test_files,
            replacements,
            create_backup=False
        )

        summary = self.batch_processor.get_summary()

        self.assertIn('total_files', summary)
        self.assertIn('successful', summary)
        self.assertIn('failed', summary)
        self.assertIn('total_replacements', summary)
        self.assertIn('success_rate', summary)

        self.assertEqual(summary['total_files'], 5)
        self.assertEqual(summary['successful'], 5)
        self.assertEqual(summary['failed'], 0)

    def test_get_files_from_directory(self):
        """Test getting files from directory."""
        # Create nested directory structure
        nested_dir = os.path.join(self.temp_dir, "nested")
        os.makedirs(nested_dir)

        nested_file = os.path.join(nested_dir, "nested.docx")
        self._create_test_document(nested_file, "Nested file 2024")

        # Get all files recursively
        files = BatchProcessor.get_files_from_directory(self.temp_dir, recursive=True)
        self.assertEqual(len(files), 6)  # 5 original + 1 nested

        # Get files non-recursively
        files = BatchProcessor.get_files_from_directory(self.temp_dir, recursive=False)
        self.assertEqual(len(files), 5)  # Only top-level


class TestProcessingResult(unittest.TestCase):
    """Test processing result class."""

    def test_successful_result(self):
        """Test successful processing result."""
        result = ProcessingResult(
            file_path="test.docx",
            success=True,
            message="Processed",
            replacements=5,
            backup_path="test_backup.docx"
        )

        result_dict = result.to_dict()
        self.assertTrue(result_dict['success'])
        self.assertEqual(result_dict['replacements'], 5)
        self.assertIn('backup_path', result_dict)

    def test_failed_result(self):
        """Test failed processing result."""
        result = ProcessingResult(
            file_path="test.docx",
            success=False,
            message="Error occurred"
        )

        result_dict = result.to_dict()
        self.assertFalse(result_dict['success'])
        self.assertEqual(result_dict['message'], "Error occurred")


def run_tests():
    """Run all tests."""
    # Create test suite
    loader = unittest.TestLoader()
    suite = unittest.TestSuite()

    # Add all test cases
    suite.addTests(loader.loadTestsFromTestCase(TestFormatPreserver))
    suite.addTests(loader.loadTestsFromTestCase(TestDocxProcessor))
    suite.addTests(loader.loadTestsFromTestCase(TestBatchProcessor))
    suite.addTests(loader.loadTestsFromTestCase(TestProcessingResult))

    # Run tests
    runner = unittest.TextTestRunner(verbosity=2)
    result = runner.run(suite)

    # Return exit code
    return 0 if result.wasSuccessful() else 1


if __name__ == '__main__':
    sys.exit(run_tests())
