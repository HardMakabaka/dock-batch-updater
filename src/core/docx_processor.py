"""Core DOCX document processor with format preservation.

This module provides the main functionality for processing DOCX documents,
including text replacement, table content modification, and strict format
preservation.
"""

import os
import shutil
import uuid
import hashlib
import re
from typing import List, Dict, Any, Optional, Tuple, Callable
from pathlib import Path
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table, _Cell
from docx.oxml.text.paragraph import CT_P

from src.utils.format_preserver import FormatPreserver


class DocxProcessor:
    """Process DOCX documents with format-preserving text replacement.

    This class provides methods to replace text in DOCX documents while
    preserving all formatting attributes including fonts, colors, sizes,
    paragraph styles, and table structures.
    """

    def __init__(self, doc_path: str):
        """Initialize the processor with a document.

        Args:
            doc_path: Path to the DOCX file
        """
        self.doc_path = doc_path
        self.doc: Optional[Document] = None
        self.backup_path: Optional[str] = None
        self.format_preserver = FormatPreserver()
        self.replacement_count = 0

    def load(self) -> bool:
        """Load the DOCX document.

        Returns:
            True if successful, False otherwise
        """
        try:
            self.doc = Document(self.doc_path)
            return True
        except Exception as e:
            print(f"Error loading document {self.doc_path}: {e}")
            return False

    def create_backup(self, backup_dir: Optional[str] = None) -> str:
        """Create a backup of the original document.

        Naming goals:
        - Uniqueness under concurrency (global backup_dir mode)
        - Human traceability (include parent dir hint + short hash)
        - Keep names short to avoid Windows path length issues

        Args:
            backup_dir: Optional directory for backups. Defaults to same directory.

        Returns:
            Path to the backup file
        """
        if backup_dir is None:
            backup_dir = os.path.dirname(self.doc_path)

        os.makedirs(backup_dir, exist_ok=True)

        original_name = os.path.basename(self.doc_path)
        name, ext = os.path.splitext(original_name)

        parent_dir = os.path.basename(os.path.dirname(self.doc_path)) or "_root_"
        parent_dir = re.sub(r"[^0-9A-Za-z_\u4e00-\u9fff]+", "_", parent_dir).strip("_")
        if not parent_dir:
            parent_dir = "_root_"
        parent_dir = parent_dir[:20]

        path_hash6 = hashlib.sha256(os.path.abspath(self.doc_path).encode("utf-8")).hexdigest()[:6]

        # Two-tier uniqueness: short uuid in final name + atomic finalization via rename.
        for _ in range(50):
            uid8 = uuid.uuid4().hex[:8]
            backup_name = f"{name}_backup_{parent_dir}_{path_hash6}_{uid8}{ext}"
            backup_path = os.path.join(backup_dir, backup_name)

            # Write to a temp file in the same directory, then move into place.
            tmp_name = f".{backup_name}.tmp"
            tmp_path = os.path.join(backup_dir, tmp_name)

            try:
                if os.path.exists(tmp_path):
                    try:
                        os.remove(tmp_path)
                    except OSError:
                        pass

                shutil.copy2(self.doc_path, tmp_path)

                # Prefer atomic replacement when available; on Windows this is also atomic.
                try:
                    os.replace(tmp_path, backup_path)
                except FileExistsError:
                    # Extremely unlikely due to uuid, but treat as collision and retry.
                    try:
                        os.remove(tmp_path)
                    except OSError:
                        pass
                    continue

                self.backup_path = backup_path
                return backup_path
            finally:
                # Best-effort cleanup if anything failed after creating tmp
                if os.path.exists(tmp_path):
                    try:
                        os.remove(tmp_path)
                    except OSError:
                        pass

        raise RuntimeError(f"Failed to create unique backup in {backup_dir}")

    def replace_text(
        self,
        search_text: str,
        replace_text: str,
        progress_callback: Optional[Callable[[int, int], None]] = None
    ) -> int:
        """Replace all occurrences of text in the document.

        This method replaces text in paragraphs and tables while preserving
        all formatting attributes.

        Args:
            search_text: Text to search for
            replace_text: Text to replace with
            progress_callback: Optional callback for progress updates

        Returns:
            Number of replacements made
        """
        if not self.doc:
            return 0

        self.replacement_count = 0
        total_items = len(self.doc.paragraphs) + len(self.doc.tables)
        processed_items = 0

        # Process paragraphs
        for paragraph in self.doc.paragraphs:
            count = self._replace_in_paragraph(paragraph, search_text, replace_text)
            self.replacement_count += count
            processed_items += 1
            if progress_callback:
                progress_callback(processed_items, total_items)

        # Process tables
        for table in self.doc.tables:
            count = self._replace_in_table(table, search_text, replace_text)
            self.replacement_count += count
            processed_items += 1
            if progress_callback:
                progress_callback(processed_items, total_items)

        return self.replacement_count

    def replace_multiple(
        self,
        replacements: List[Tuple[str, str]],
        progress_callback: Optional[Callable[[int, int], None]] = None
    ) -> int:
        """Replace multiple text patterns in the document.

        Args:
            replacements: List of (search_text, replace_text) tuples
            progress_callback: Optional callback for progress updates

        Returns:
            Total number of replacements made
        """
        if not self.doc:
            return 0

        total_replacements = 0

        for search_text, replace_text in replacements:
            count = self.replace_text(search_text, replace_text, progress_callback)
            total_replacements += count

        return total_replacements

    def _replace_in_paragraph(
        self,
        paragraph: Paragraph,
        search_text: str,
        replace_text: str
    ) -> int:
        """Replace text in a paragraph while preserving format.

        Args:
            paragraph: A python-docx Paragraph object
            search_text: Text to search for
            replace_text: Text to replace with

        Returns:
            Number of replacements made
        """
        if search_text not in paragraph.text:
            return 0

        count = 0
        search_start = 0

        while True:
            paragraph_text = paragraph.text
            match_index = paragraph_text.find(search_text, search_start)
            if match_index == -1:
                break

            start_pos = match_index
            end_pos = match_index + len(search_text)

            current_pos = 0
            start_run_idx = None
            end_run_idx = None
            start_offset = 0
            end_offset = 0

            for run_idx, run in enumerate(paragraph.runs):
                run_text = run.text
                run_end = current_pos + len(run_text)

                if start_run_idx is None and start_pos < run_end:
                    start_run_idx = run_idx
                    start_offset = start_pos - current_pos

                if end_pos <= run_end:
                    end_run_idx = run_idx
                    end_offset = end_pos - current_pos
                    break

                current_pos = run_end

            if start_run_idx is None or end_run_idx is None:
                break

            start_run = paragraph.runs[start_run_idx]
            end_run = paragraph.runs[end_run_idx]
            format_data = self.format_preserver.capture_run_format(start_run)

            new_text = (
                start_run.text[:start_offset]
                + replace_text
                + end_run.text[end_offset:]
            )
            start_run.text = new_text
            self.format_preserver.apply_run_format(start_run, format_data)

            for idx in range(start_run_idx + 1, end_run_idx + 1):
                paragraph.runs[idx].text = ""

            count += 1
            search_start = match_index + len(replace_text)

        return count

    def _replace_in_table(
        self,
        table: Table,
        search_text: str,
        replace_text: str
    ) -> int:
        """Replace text in a table while preserving format.

        Args:
            table: A python-docx Table object
            search_text: Text to search for
            replace_text: Text to replace with

        Returns:
            Number of replacements made
        """
        count = 0

        for row in table.rows:
            for cell in row.cells:
                # Process paragraphs within the cell
                for paragraph in cell.paragraphs:
                    cell_count = self._replace_in_paragraph(paragraph, search_text, replace_text)
                    count += cell_count

                # Process nested tables
                for nested_table in cell.tables:
                    nested_count = self._replace_in_table(nested_table, search_text, replace_text)
                    count += nested_count

        return count

    def save(self, output_path: Optional[str] = None) -> bool:
        """Save the document.

        Args:
            output_path: Optional path to save to. Defaults to original path.

        Returns:
            True if successful, False otherwise
        """
        if not self.doc:
            return False

        try:
            save_path = output_path or self.doc_path
            self.doc.save(save_path)
            return True
        except Exception as e:
            print(f"Error saving document: {e}")
            return False

    def get_statistics(self) -> Dict[str, Any]:
        """Get statistics about the document.

        Returns:
            Dictionary containing document statistics
        """
        if not self.doc:
            return {}

        stats = {
            'paragraph_count': len(self.doc.paragraphs),
            'table_count': len(self.doc.tables),
            'character_count': sum(len(p.text) for p in self.doc.paragraphs),
            'word_count': sum(len(p.text.split()) for p in self.doc.paragraphs),
            'replacement_count': self.replacement_count
        }

        # Count cells in all tables
        total_cells = 0
        for table in self.doc.tables:
            for row in table.rows:
                total_cells += len(row.cells)
        stats['cell_count'] = total_cells

        return stats

    def validate_document(self) -> Tuple[bool, List[str]]:
        """Validate that the document is properly loaded.

        Returns:
            Tuple of (is_valid, list_of_errors)
        """
        errors = []

        if not self.doc:
            errors.append("Document not loaded")
            return False, errors

        try:
            # Try to access document properties
            _ = self.doc.paragraphs
            _ = self.doc.tables
        except Exception as e:
            errors.append(f"Document structure error: {str(e)}")

        return len(errors) == 0, errors

    def restore_backup(self) -> bool:
        """Restore the document from backup.

        Returns:
            True if successful, False otherwise
        """
        if not self.backup_path or not os.path.exists(self.backup_path):
            return False

        try:
            shutil.copy2(self.backup_path, self.doc_path)
            # Reload the document
            return self.load()
        except Exception as e:
            print(f"Error restoring backup: {e}")
            return False

    @staticmethod
    def is_docx_file(file_path: str) -> bool:
        """Check if a file is a valid DOCX file.

        Args:
            file_path: Path to the file

        Returns:
            True if the file is a DOCX file
        """
        # Check extension
        if not file_path.lower().endswith('.docx'):
            return False

        # Check if file exists
        if not os.path.exists(file_path):
            return False

        # Try to open as a ZIP file (DOCX is a ZIP archive)
        try:
            import zipfile
            with zipfile.ZipFile(file_path, 'r') as zip_ref:
                # Check for required DOCX files
                required_files = ['[Content_Types].xml', '_rels/.rels', 'word/document.xml']
                for req_file in required_files:
                    if req_file not in zip_ref.namelist():
                        return False
            return True
        except Exception:
            return False

    def close(self) -> None:
        """Clean up resources."""
        self.doc = None
