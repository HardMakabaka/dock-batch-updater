"""
补充 DocxProcessor 的边界与异常分支测试，提升覆盖率。
"""

import os
import sys
import tempfile
import shutil
import unittest

# 添加 src 到路径以便导入
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'src'))

try:
    from docx import Document
    from core.docx_processor import DocxProcessor
except ImportError as e:
    print(f"Import error: {e}")
    print("Please install dependencies: pip install -r requirements.txt")
    sys.exit(1)


class TestDocxProcessorAdditional(unittest.TestCase):
    """DocxProcessor 边界条件与异常分支测试。"""

    def setUp(self):
        self.temp_dir = tempfile.mkdtemp()
        self.test_file = os.path.join(self.temp_dir, "sample.docx")

        doc = Document()
        doc.add_paragraph("Hello 2024")
        doc.add_paragraph("Another 2024")
        table = doc.add_table(rows=1, cols=1)
        table.rows[0].cells[0].text = "Cell 2024"
        doc.save(self.test_file)

    def tearDown(self):
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)

    def test_load_failure(self):
        """加载不存在文件应返回 False。"""
        processor = DocxProcessor(os.path.join(self.temp_dir, "missing.docx"))
        self.assertFalse(processor.load())

    def test_replace_text_without_load(self):
        """未加载文档时替换应返回 0。"""
        processor = DocxProcessor(self.test_file)
        self.assertEqual(processor.replace_text("2024", "2025"), 0)

    def test_replace_multiple(self):
        """多规则替换应返回累计数量。"""
        processor = DocxProcessor(self.test_file)
        processor.load()

        count = processor.replace_multiple([
            ("2024", "2025"),
            ("Hello", "Hi")
        ])

        self.assertGreaterEqual(count, 4)

    def test_replace_multiple_without_load(self):
        """未加载文档时多规则替换应返回 0。"""
        processor = DocxProcessor(self.test_file)
        self.assertEqual(processor.replace_multiple([("2024", "2025")]), 0)

    def test_replace_text_progress_callback(self):
        """替换时进度回调应按段落与表格数量触发。"""
        processor = DocxProcessor(self.test_file)
        processor.load()
        calls = []

        def progress_callback(done, total):
            calls.append((done, total))

        processor.replace_text("2024", "2025", progress_callback)
        self.assertEqual(len(calls), 3)

    def test_save_without_doc(self):
        """未加载文档时保存应返回 False。"""
        processor = DocxProcessor(self.test_file)
        self.assertFalse(processor.save())

    def test_save_failure_invalid_path(self):
        """保存到不存在目录应返回 False。"""
        processor = DocxProcessor(self.test_file)
        processor.load()

        invalid_path = os.path.join(self.temp_dir, "not_exists", "out.docx")
        self.assertFalse(processor.save(invalid_path))

    def test_validate_document_not_loaded(self):
        """未加载文档时校验应失败。"""
        processor = DocxProcessor(self.test_file)
        is_valid, errors = processor.validate_document()

        self.assertFalse(is_valid)
        self.assertIn("Document not loaded", errors)

    def test_validate_document_structure_error(self):
        """文档结构异常时应返回错误信息。"""
        processor = DocxProcessor(self.test_file)

        class BrokenDoc:
            @property
            def paragraphs(self):
                raise RuntimeError("broken paragraphs")

            @property
            def tables(self):
                raise RuntimeError("broken tables")

        processor.doc = BrokenDoc()
        is_valid, errors = processor.validate_document()

        self.assertFalse(is_valid)
        self.assertTrue(any("Document structure error" in msg for msg in errors))

    def test_validate_document_loaded(self):
        """加载后校验应通过。"""
        processor = DocxProcessor(self.test_file)
        processor.load()
        is_valid, errors = processor.validate_document()

        self.assertTrue(is_valid)
        self.assertEqual(errors, [])

    def test_restore_backup_no_backup(self):
        """未创建备份时恢复应返回 False。"""
        processor = DocxProcessor(self.test_file)
        processor.load()
        self.assertFalse(processor.restore_backup())

    def test_restore_backup_success(self):
        """创建备份后应可恢复。"""
        processor = DocxProcessor(self.test_file)
        processor.load()
        backup_path = processor.create_backup()

        # 修改原文件内容
        doc = Document()
        doc.add_paragraph("Changed")
        doc.save(self.test_file)

        self.assertTrue(os.path.exists(backup_path))
        self.assertTrue(processor.restore_backup())

    def test_restore_backup_exception(self):
        """恢复过程中出现异常应返回 False。"""
        processor = DocxProcessor(self.test_file)
        processor.load()

        backup_path = processor.create_backup()
        processor.backup_path = backup_path
        processor.doc_path = os.path.join(self.temp_dir, "not_exists", "restore.docx")

        self.assertFalse(processor.restore_backup())

    def test_create_backup_duplicate(self):
        """重复备份应生成不同文件名的备份。"""
        processor = DocxProcessor(self.test_file)
        processor.load()
        first_backup = processor.create_backup()
        second_backup = processor.create_backup()

        self.assertTrue(os.path.exists(first_backup))
        self.assertTrue(os.path.exists(second_backup))
        self.assertNotEqual(first_backup, second_backup)

    def test_get_statistics_without_load(self):
        """未加载文档时统计应返回空字典。"""
        processor = DocxProcessor(self.test_file)
        self.assertEqual(processor.get_statistics(), {})

    def test_replace_text_with_nested_table(self):
        """嵌套表格内容应被替换。"""
        nested_file = os.path.join(self.temp_dir, "nested.docx")
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        nested_table = table.rows[0].cells[0].add_table(rows=1, cols=1)
        nested_table.rows[0].cells[0].text = "Nested 2024"
        doc.save(nested_file)

        processor = DocxProcessor(nested_file)
        processor.load()
        count = processor.replace_text("2024", "2025")

        self.assertEqual(count, 1)

    def test_replace_text_across_runs(self):
        """跨多个 run 的文本应被替换。"""
        run_file = os.path.join(self.temp_dir, "runs.docx")
        doc = Document()
        paragraph = doc.add_paragraph()
        paragraph.add_run("Hello")
        paragraph.add_run(" World")
        doc.save(run_file)

        processor = DocxProcessor(run_file)
        processor.load()
        count = processor.replace_text("Hello World", "Hi")

        self.assertEqual(count, 1)
        self.assertEqual(processor.doc.paragraphs[0].text, "Hi")

    def test_close_clears_doc(self):
        """关闭后 doc 应为 None。"""
        processor = DocxProcessor(self.test_file)
        processor.load()
        processor.close()

        self.assertIsNone(processor.doc)


if __name__ == "__main__":
    unittest.main(verbosity=2)
